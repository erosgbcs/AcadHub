import json
import io
import zipfile
import xml.etree.ElementTree as ET
import re
import random
import time
import asyncio
import hashlib
from collections import Counter
from typing import List, Dict, Optional, Any

from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field
import aiohttp

# PDF Processing Libraries
import pypdf

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except Exception:
    PDFPLUMBER_AVAILABLE = False

# Word Document Processing (.docx)
try:
    import docx
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# spaCy for High-Accuracy NLP
try:
    import spacy
    NLP = spacy.load("en_core_web_sm")
except Exception:
    NLP = None
    print("Warning: spaCy model not available. Falling back to regex heuristics.")

# scikit-learn for TF-IDF TextRank Summarization
try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    SKLEARN_AVAILABLE = True
except Exception:
    SKLEARN_AVAILABLE = False

# ------------------------------
# App Initialization
# ------------------------------
app = FastAPI(title="AcademicHub Engine - High Accuracy Reviewer Generator")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ------------------------------
# Pydantic Models
# ------------------------------
class ReviewRequest(BaseModel):
    notes: str = ""
    num_flashcards: int = 12
    quiz_types: Dict[str, int] = {"identification": 5}
    use_internet: bool = False
    enrich_count: int = 5

# ------------------------------
# Utility Constants
# ------------------------------
SENTENCE_PATTERN = re.compile(r'(?<!\d)[.!?](?!\d)')
WORD_PATTERN = re.compile(r'[a-zA-Z]+')
CAPITALIZED_PHRASE = re.compile(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b')
TECH_TERM = re.compile(r'\b[A-Za-z]+[#+]\b')
STOP_WORDS = frozenset({
    'the', 'is', 'at', 'which', 'on', 'a', 'an', 'and', 'or', 'but', 'in',
    'with', 'to', 'for', 'of', 'from', 'by', 'as', 'be', 'was', 'are', 'been',
    'this', 'that', 'these', 'those', 'it', 'its', 'they', 'them', 'we', 'you',
    'he', 'she', 'his', 'her', 'their', 'our', 'my', 'your', 'has', 'have',
    'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may',
    'might', 'can', 'shall', 'not', 'no', 'so', 'if', 'then', 'than', 'too',
    'very', 'just', 'about', 'also', 'into', 'onto', 'upon', 'within',
    'without', 'because', 'each', 'all', 'some', 'any', 'every', 'both',
    'few', 'more', 'most', 'other', 'such', 'only', 'own', 'same', 'new'
})

# ------------------------------
# Advanced Document Text Extractors
# ------------------------------
def extract_pdf_text(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber if available, falling back to pypdf."""
    extracted_text = []
    
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        extracted_text.append(text)
            if extracted_text:
                return "\n".join(extracted_text)
        except Exception:
            pass

    # Fallback to pypdf
    try:
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            text = page.extract_text()
            if text:
                extracted_text.append(text)
        return "\n".join(extracted_text)
    except Exception:
        raise HTTPException(400, "Error reading PDF file. Ensure file is not corrupt or password protected.")

def extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from Word (.docx) documents including paragraphs and tables."""
    if DOCX_AVAILABLE:
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            full_text = []
            
            # Extract Paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            # Extract Tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))
                        
            return "\n".join(full_text)
        except Exception:
            pass

    # XML Zip Fallback
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            with z.open("word/document.xml") as f:
                tree = ET.parse(f)
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                paragraphs = []
                for p in tree.iterfind('.//w:p', ns):
                    texts = [t.text for t in p.iterfind('.//w:t', ns) if t.text]
                    if texts:
                        paragraphs.append(''.join(texts))
                return '\n'.join(paragraphs)
    except Exception:
        raise HTTPException(400, "DOCX read error.")

def extract_rtf_text(file_bytes: bytes) -> str:
    try:
        text = file_bytes.decode('utf-8', errors='ignore')
        text = re.sub(r'\\[a-z]+\-?\d* ?', '', text)
        text = re.sub(r'[{}]', '', text)
        text = re.sub(r'\\\'[0-9a-fA-F]{2}', ' ', text)
        return text
    except Exception:
        raise HTTPException(400, "RTF read error.")

def extract_html_text(file_bytes: bytes) -> str:
    try:
        html = file_bytes.decode('utf-8', errors='ignore')
        html = re.sub(r'<script[^>]*>.*?</script>', ' ', html, flags=re.DOTALL)
        html = re.sub(r'<style[^>]*>.*?</style>', ' ', html, flags=re.DOTALL)
        html = re.sub(r'<[^>]+>', ' ', html)
        html = html.replace('&nbsp;', ' ').replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
        return html
    except Exception:
        raise HTTPException(400, "HTML read error.")

def extract_document_text(file_bytes: bytes, filename: str) -> str:
    """Unified Document Router."""
    fname = filename.lower()
    if fname.endswith(".pdf"):
        return extract_pdf_text(file_bytes)
    elif fname.endswith((".docx", ".doc")):
        return extract_docx_text(file_bytes)
    elif fname.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="ignore")
    elif fname.endswith(".md"):
        return file_bytes.decode("utf-8", errors="ignore")
    elif fname.endswith(".rtf"):
        return extract_rtf_text(file_bytes)
    elif fname.endswith((".html", ".htm")):
        return extract_html_text(file_bytes)
    else:
        raise HTTPException(400, "Unsupported document type. Please upload PDF, DOCX, TXT, MD, RTF, or HTML.")

# ------------------------------
# High-Accuracy NLP Engine
# ------------------------------
def smart_sentences(text: str) -> List[str]:
    if NLP is not None:
        doc = NLP(text)
        return [sent.text.strip() for sent in doc.sents if len(sent.text.strip()) > 20]
    
    text = re.sub(r'(Dr|Mr|Mrs|Ms|Prof|etc|vs|e\.g|i\.e)\.', r'\1<DOT>', text)
    sents = SENTENCE_PATTERN.split(text)
    return [s.strip().replace('<DOT>', '.') for s in sents if len(s.strip()) > 20]

def extract_key_phrases(text: str, top_n: int = 20) -> List[str]:
    if NLP is not None:
        doc = NLP(text)
        phrases = []
        for chunk in doc.noun_chunks:
            words = [token.text for token in chunk if not token.is_stop and not token.is_punct]
            if len(words) >= 1 and (len(words) >= 2 or chunk.root.pos_ in ("PROPN", "NOUN")):
                phrase = " ".join(words).strip()
                if phrase and phrase.lower() not in STOP_WORDS and len(phrase) > 2:
                    phrases.append(phrase)
        
        for ent in doc.ents:
            if ent.label_ in ("PERSON", "ORG", "PRODUCT", "GPE", "EVENT", "WORK_OF_ART", "TECHNOLOGY"):
                if ent.text.lower() not in STOP_WORDS:
                    phrases.append(ent.text.strip())

        phrase_counts = Counter(phrases)
        weighted = [(p, c * (len(p.split()) ** 1.5)) for p, c in phrase_counts.items()]
        weighted.sort(key=lambda x: x[1], reverse=True)
        
        final = []
        for phrase, _ in weighted:
            if not any(phrase != other and phrase.lower() in other.lower() for other in final):
                final.append(phrase)
            if len(final) >= top_n:
                break
        return final

    phrases = CAPITALIZED_PHRASE.findall(text) + TECH_TERM.findall(text)
    phrase_counts = Counter([p for p in phrases if p.lower() not in STOP_WORDS])
    return [p[0] for p in phrase_counts.most_common(top_n)]

def find_definition_sentence(text: str, phrase: str) -> str:
    sents = smart_sentences(text)
    phrase_lower = phrase.lower()
    
    scored_sents = []
    for sent in sents:
        sent_lower = sent.lower()
        if phrase_lower not in sent_lower:
            continue
        
        score = 0
        if re.search(r'\b(?:is|are|was|were)\s+(?:a|an|the|defined as|known as|a type of)\b', sent_lower):
            score += 5
        if re.search(r'\b(?:refers to|means|denotes|describes|consists of)\b', sent_lower):
            score += 5
        if '(' in sent and ')' in sent:
            score += 2
        if sent_lower.startswith(phrase_lower):
            score += 3
        
        scored_sents.append((sent, score))
    
    if scored_sents:
        scored_sents.sort(key=lambda x: x[1], reverse=True)
        return scored_sents[0][0][:300].strip()
        
    return f"{phrase} is a key topic covered in the document."

def summarize_text(text: str, max_sentences: int = 5) -> List[str]:
    sents = smart_sentences(text)
    if len(sents) <= max_sentences:
        return sents

    if SKLEARN_AVAILABLE:
        vectorizer = TfidfVectorizer(stop_words='english')
        try:
            tfidf_matrix = vectorizer.fit_transform(sents)
            similarity_matrix = cosine_similarity(tfidf_matrix, tfidf_matrix)
            scores = similarity_matrix.sum(axis=1)
            ranked = sorted(range(len(sents)), key=lambda i: scores[i], reverse=True)
            top_indices = sorted(ranked[:max_sentences])
            return [sents[i] for i in top_indices]
        except Exception:
            pass

    return sents[:max_sentences]

def get_smart_distractors(correct: str, pool: List[str], count: int = 3) -> List[str]:
    target_len = len(correct)
    filtered = [
        p for p in pool 
        if p.lower() != correct.lower() and (0.6 * target_len <= len(p) <= 1.4 * target_len)
    ]
    if len(filtered) < count:
        filtered = [p for p in pool if p.lower() != correct.lower()]
        
    if len(filtered) >= count:
        return random.sample(filtered, count)
    
    distractors = list(filtered)
    generic = ["None of the above", "All of the above", "Not specified in lesson", "Other related factors"]
    for g in generic:
        if len(distractors) >= count:
            break
        if g not in distractors:
            distractors.append(g)
            
    return distractors[:count]

def deduplicate_options(correct: str, raw_distractors: List[str]) -> List[str]:
    options = [correct]
    for d in raw_distractors:
        if d not in options and d.strip():
            options.append(d)
    while len(options) < 4:
        options.append(f"Option {len(options) + 1}")
    return options[:4]

# ------------------------------
# Core Generation Pipeline
# ------------------------------
def generate_reviewer_content(text: str, num_flashcards: int, quiz_types: Dict[str, int]):
    if not text.strip():
        raise ValueError("No text extracted from document.")

    start = time.time()
    sents = smart_sentences(text)
    phrases = extract_key_phrases(text, top_n=max(num_flashcards, 30))
    doc = NLP(text) if NLP is not None else None

    summary = summarize_text(text, max_sentences=5)

    # Generate Flashcards
    flashcards = []
    used_defs = set()
    for phrase in phrases:
        if len(flashcards) >= num_flashcards:
            break
        def_sent = find_definition_sentence(text, phrase)
        if def_sent in used_defs:
            continue
        flashcards.append({
            "id": len(flashcards) + 1,
            "term": phrase.strip(),
            "definition": def_sent
        })
        used_defs.add(def_sent)

    # Generate Quizzes
    quiz = []

    # True/False
    tf_count = quiz_types.get("truefalse", 0)
    for i in range(tf_count):
        sent = random.choice(sents) if sents else ""
        if not sent:
            break
        
        is_false_target = (i % 2 == 1)
        q_text = sent[:200]
        correct = "True"

        if is_false_target and doc is not None:
            sent_doc = NLP(sent)
            if sent_doc.ents:
                target_ent = random.choice(sent_doc.ents)
                same_label_ents = [e.text for e in doc.ents if e.label_ == target_ent.label_ and e.text.lower() != target_ent.text.lower()]
                if same_label_ents:
                    false_val = random.choice(same_label_ents)
                    q_text = sent.replace(target_ent.text, false_val)[:200]
                    correct = "False"

        quiz.append({
            "id": len(quiz) + 1,
            "type": "truefalse",
            "question": f'True or False: "{q_text}"',
            "options": ["True", "False"],
            "answer": correct,
            "explanation": f"The statement is {correct.lower()}."
        })

    # Identification
    id_count = quiz_types.get("identification", 0)
    for _ in range(id_count):
        if not phrases:
            break
        phrase = random.choice(phrases)
        def_sent = find_definition_sentence(text, phrase)
        blanked = re.sub(re.escape(phrase), "________", def_sent, flags=re.IGNORECASE)
        
        distractor_pool = [p for p in phrases if p != phrase]
        distractors = get_smart_distractors(phrase, distractor_pool, 3)
        options = deduplicate_options(phrase, distractors)
        random.shuffle(options)

        quiz.append({
            "id": len(quiz) + 1,
            "type": "identification",
            "question": f'Fill in the blank: "{blanked}"',
            "options": options,
            "answer": phrase,
            "explanation": f"The correct term is '{phrase}'."
        })

    # Multiple Choice
    mc_count = quiz_types.get("multiplechoice", 0)
    for _ in range(mc_count):
        if not phrases:
            break
        phrase = random.choice(phrases)
        correct_def = find_definition_sentence(text, phrase)
        
        other_defs = [find_definition_sentence(text, p) for p in phrases if p != phrase]
        distractors = get_smart_distractors(correct_def, other_defs, 3)
        
        options = deduplicate_options(correct_def, distractors)
        random.shuffle(options)

        quiz.append({
            "id": len(quiz) + 1,
            "type": "multiplechoice",
            "question": f"What is '{phrase}'?",
            "options": options,
            "answer": correct_def,
            "explanation": f"Definition of {phrase}."
        })

    elapsed = round(time.time() - start, 2)
    return {
        "summary": summary,
        "flashcards": flashcards,
        "quiz": quiz,
        "metadata": {
            "method": "accuracy_v7_document_nlp",
            "length": len(text),
            "time": elapsed,
            "num_sentences": len(sents),
            "num_phrases": len(phrases)
        }
    }

# ------------------------------
# API Endpoints
# ------------------------------
@app.get("/")
def root():
    return {"status": "online", "system": "AcademicHub High-Precision Document Engine"}

@app.get("/api/health")
def health():
    return {
        "status": "healthy",
        "nlp": "spacy" if NLP else "fallback",
        "sklearn": SKLEARN_AVAILABLE,
        "pdfplumber": PDFPLUMBER_AVAILABLE,
        "python-docx": DOCX_AVAILABLE
    }

@app.post("/api/generate-reviewer-local")
async def gen_local(
    notes: str = Form(""),
    file: UploadFile = File(None),
    num_flashcards: int = Form(12),
    quiz_types: str = Form('{"identification":5}')
):
    text = notes.strip() if notes else ""
    if file:
        fb = await file.read()
        extracted = extract_document_text(fb, file.filename)
        text += "\n" + extracted

    if not text.strip():
        raise HTTPException(400, "Please provide lesson text or upload a valid document (PDF, DOCX, TXT).")

    try:
        qt = json.loads(quiz_types)
    except:
        qt = {"identification": 5}

    result = generate_reviewer_content(text, num_flashcards, qt)
    return JSONResponse(content=result)

@app.post("/api/reviewer")
async def generate_reviewer_json(request: ReviewRequest):
    text = request.notes.strip()
    if not text:
        raise HTTPException(400, "No text provided.")
    result = generate_reviewer_content(text, request.num_flashcards, request.quiz_types)
    return JSONResponse(content=result)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)